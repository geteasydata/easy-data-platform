"""
Word Report Generator - Editable Word Documents
"""

from typing import Dict, Any, List, Optional
import io
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_word_report(
    analysis: Dict[str, Any],
    results: Dict[str, Any],
    target_col: str,
    insights: str,
    lang: str = 'ar'
) -> Optional[bytes]:
    """
    Create an editable Word report.
    
    Args:
        analysis: Data analysis dictionary
        results: ML results dictionary
        target_col: Target column name
        insights: AI-generated insights
        lang: Language ('ar' or 'en')
    
    Returns:
        Word document as bytes, or None if python-docx not available
    """
    if not HAS_DOCX:
        return None
    
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Expert - Data Science Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Target Variable: {target_col}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(46, 134, 171)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # Section 1: Data Summary
    section1 = '1. Data Summary' if lang == 'en' else '1. Data Summary (ملخص البيانات)'
    doc.add_heading(section1, level=1)
    
    # Summary table
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Handle nested analysis structure
    overview = analysis.get('overview', analysis)
    cols_val = overview.get('columns', analysis.get('columns', 0))
    cols_count = len(cols_val) if isinstance(cols_val, dict) else cols_val
    
    summary_data = [
        ('Rows', str(overview.get('rows', 0))),
        ('Columns', str(cols_count)),
        ('Numeric Columns', str(len(analysis.get('numeric_columns', [])))),
        ('Categorical Columns', str(len(analysis.get('categorical_columns', [])))),
        ('Missing Values', str(overview.get('total_missing', 0))),
        ('Duplicates', str(overview.get('duplicates', 0))),
        ('Memory (MB)', f"{overview.get('memory_mb', 0):.2f}"),
    ]
    
    for i, (key, value) in enumerate(summary_data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        # Style first column
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Section 2: Model Performance
    section2 = '2. Model Performance' if lang == 'en' else '2. Model Performance (أداء النموذج)'
    doc.add_heading(section2, level=1)
    
    p = doc.add_paragraph()
    p.add_run('Problem Type: ').bold = True
    p.add_run(results['problem_type'].title())
    
    p = doc.add_paragraph()
    p.add_run('Best Model: ').bold = True
    p.add_run(results['best_model'])
    
    # Metrics table
    metrics = results['metrics']
    table = doc.add_table(rows=len(metrics), cols=2)
    table.style = 'Table Grid'
    
    for i, (key, value) in enumerate(metrics.items()):
        row = table.rows[i]
        row.cells[0].text = key.replace('_', ' ').title()
        if isinstance(value, float):
            row.cells[1].text = f"{value:.4f}"
        else:
            row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Section 3: Feature Importance
    section3 = '3. Feature Importance' if lang == 'en' else '3. Feature Importance (أهمية الميزات)'
    doc.add_heading(section3, level=1)
    
    importance = results['feature_importance'].head(10)
    table = doc.add_table(rows=len(importance) + 1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    header = table.rows[0]
    header.cells[0].text = 'Rank'
    header.cells[1].text = 'Feature'
    header.cells[2].text = 'Importance'
    for cell in header.cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '2E86AB')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data
    for i, (idx, row) in enumerate(importance.iterrows()):
        table_row = table.rows[i + 1]
        table_row.cells[0].text = str(i + 1)
        table_row.cells[1].text = str(row['Feature'])
        table_row.cells[2].text = f"{row['Importance']:.4f}"
    
    doc.add_paragraph()
    
    # Section 4: AI Insights
    section4 = '4. AI Expert Insights' if lang == 'en' else '4. AI Expert Insights (رؤى الذكاء الاصطناعي)'
    doc.add_heading(section4, level=1)
    
    # Clean and add insights
    clean_insights = insights.replace('**', '').replace('*', '')
    for line in clean_insights.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    doc.add_paragraph()
    
    # Section 5: Cleaning Steps
    section5 = '5. Data Cleaning Steps' if lang == 'en' else '5. Data Cleaning Steps (خطوات التنظيف)'
    doc.add_heading(section5, level=1)
    
    for step in results['cleaning_steps']:
        clean_step = step.replace('✅', '[OK]').replace('🗑️', '[DEL]').replace('🔧', '[FIX]').replace('🏷️', '[ENC]').replace('🎯', '[TGT]').replace('✨', '[OK]')
        doc.add_paragraph(clean_step, style='List Bullet')
    
    # All Models Comparison
    doc.add_heading('6. All Models Comparison', level=1)
    
    all_models = results.get('all_models', {})
    if all_models:
        table = doc.add_table(rows=len(all_models) + 1, cols=2)
        table.style = 'Table Grid'
        
        header = table.rows[0]
        header.cells[0].text = 'Model'
        header.cells[1].text = 'Score'
        for cell in header.cells:
            cell.paragraphs[0].runs[0].bold = True
        
        sorted_models = sorted(all_models.items(), key=lambda x: x[1], reverse=True)
        for i, (model_name, score) in enumerate(sorted_models):
            row = table.rows[i + 1]
            row.cells[0].text = model_name
            row.cells[1].text = f"{score:.4f}"
            
            # Highlight best model
            if model_name == results['best_model']:
                for cell in row.cells:
                    set_cell_shading(cell, 'A3BE8C')
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    
    return buffer.getvalue()
