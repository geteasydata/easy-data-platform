"""
Report Export Script
Exports reports to Word and PDF formats
"""

import pandas as pd
import argparse
import sys
from pathlib import Path
from datetime import datetime

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from paths.data_analyst.analyzer import DataAnalyzer
from paths.data_analyst.insights import InsightsGenerator
from scripts.load_data import DataLoader


class ReportExporter:
    """Export reports to various formats"""
    
    def __init__(self, output_dir: Path = None):
        self.output_dir = output_dir or Path("outputs/reports")
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def export_word(self, df: pd.DataFrame, analysis: dict, 
                    insights: list, domain: str) -> str:
        """Export report to Word format"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return self._export_markdown(df, analysis, insights, domain)
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'{domain.upper()} Data Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(
            f"This report analyzes a dataset containing {analysis['rows']:,} records "
            f"across {analysis['columns']} variables. The analysis identified "
            f"{analysis['issues_count']} potential data quality issues and generated "
            f"{len(insights)} actionable insights."
        )
        
        # Data Overview
        doc.add_heading('Data Overview', level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        metrics = [
            ('Total Records', f"{analysis['rows']:,}"),
            ('Total Columns', str(analysis['columns'])),
            ('Missing Data', f"{analysis['missing_percent']}%"),
            ('Duplicate Rows', f"{analysis['duplicate_rows']:,}"),
            ('Strong Correlations', str(analysis['strong_correlations']))
        ]
        
        for i, (label, value) in enumerate(metrics):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        
        # Key Insights
        doc.add_heading('Key Insights', level=1)
        for i, insight in enumerate(insights[:10], 1):
            doc.add_heading(f"{i}. {insight['title']}", level=2)
            doc.add_paragraph(insight['description'])
            if insight.get('recommendation'):
                rec = doc.add_paragraph()
                rec.add_run('Recommendation: ').bold = True
                rec.add_run(insight['recommendation'])
        
        # Save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = self.output_dir / f"report_{domain}_{timestamp}.docx"
        doc.save(filename)
        
        return str(filename)
    
    def _export_markdown(self, df: pd.DataFrame, analysis: dict,
                         insights: list, domain: str) -> str:
        """Export report as Markdown (fallback)"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = self.output_dir / f"report_{domain}_{timestamp}.md"
        
        content = f"""# {domain.upper()} Data Analysis Report
*Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}*

## Executive Summary
This report analyzes a dataset containing {analysis['rows']:,} records 
across {analysis['columns']} variables. The analysis identified 
{analysis['issues_count']} potential data quality issues and generated 
{len(insights)} actionable insights.

## Data Overview
| Metric | Value |
|--------|-------|
| Total Records | {analysis['rows']:,} |
| Total Columns | {analysis['columns']} |
| Missing Data | {analysis['missing_percent']}% |
| Duplicate Rows | {analysis['duplicate_rows']:,} |
| Strong Correlations | {analysis['strong_correlations']} |

## Key Insights
"""
        
        for i, insight in enumerate(insights[:10], 1):
            content += f"\n### {i}. {insight['title']}\n"
            content += f"{insight['description']}\n"
            if insight.get('recommendation'):
                content += f"\n**Recommendation:** {insight['recommendation']}\n"
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return str(filename)
    
    def export_pdf(self, df: pd.DataFrame, analysis: dict,
                   insights: list, domain: str) -> str:
        """Export report to PDF format"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        except ImportError:
            # Fallback to markdown
            md_file = self._export_markdown(df, analysis, insights, domain)
            return f"PDF export requires reportlab. Markdown saved to: {md_file}"
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = self.output_dir / f"report_{domain}_{timestamp}.pdf"
        
        doc = SimpleDocTemplate(str(filename), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph(f"{domain.upper()} Data Analysis Report", title_style))
        story.append(Spacer(1, 12))
        
        # Date
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 24))
        
        # Summary
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        summary_text = (
            f"This report analyzes a dataset containing {analysis['rows']:,} records "
            f"across {analysis['columns']} variables. The analysis identified "
            f"{analysis['issues_count']} potential data quality issues."
        )
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 24))
        
        # Data Overview Table
        story.append(Paragraph("Data Overview", styles['Heading2']))
        table_data = [
            ['Metric', 'Value'],
            ['Total Records', f"{analysis['rows']:,}"],
            ['Total Columns', str(analysis['columns'])],
            ['Missing Data', f"{analysis['missing_percent']}%"],
            ['Duplicate Rows', f"{analysis['duplicate_rows']:,}"]
        ]
        
        t = Table(table_data, colWidths=[200, 200])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(t)
        story.append(Spacer(1, 24))
        
        # Insights
        story.append(Paragraph("Key Insights", styles['Heading2']))
        for i, insight in enumerate(insights[:5], 1):
            story.append(Paragraph(f"{i}. {insight['title']}", styles['Heading3']))
            story.append(Paragraph(insight['description'], styles['Normal']))
            story.append(Spacer(1, 12))
        
        doc.build(story)
        return str(filename)


def export_report(input_file: str, format: str = 'all',
                  domain: str = 'custom', output_dir: str = None):
    """
    Export analysis report
    
    Args:
        input_file: Path to input data file
        format: Output format ('word', 'pdf', 'markdown', 'all')
        domain: Business domain
        output_dir: Output directory
    """
    # Load and analyze data
    loader = DataLoader()
    df = loader.load(input_file)
    
    analyzer = DataAnalyzer()
    report = analyzer.analyze(df)
    analysis = analyzer.get_summary()
    
    insights_gen = InsightsGenerator(domain=domain)
    insights = insights_gen.generate_insights(df, analysis)
    insights_dict = insights_gen.to_dict()
    
    # Set output directory
    output_path = Path(output_dir) if output_dir else None
    exporter = ReportExporter(output_dir=output_path)
    
    print(f"\n📄 Exporting reports...")
    
    if format in ['word', 'all']:
        path = exporter.export_word(df, analysis, insights_dict, domain)
        print(f"   ✅ Word: {path}")
    
    if format in ['pdf', 'all']:
        path = exporter.export_pdf(df, analysis, insights_dict, domain)
        print(f"   ✅ PDF: {path}")
    
    if format in ['markdown', 'all']:
        path = exporter._export_markdown(df, analysis, insights_dict, domain)
        print(f"   ✅ Markdown: {path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Export analysis report')
    parser.add_argument('--file', '-f', required=True, help='Input data file')
    parser.add_argument('--format', '-fmt', default='all',
                        choices=['word', 'pdf', 'markdown', 'all'],
                        help='Report format')
    parser.add_argument('--domain', '-d', default='custom', help='Business domain')
    parser.add_argument('--output', '-o', help='Output directory')
    
    args = parser.parse_args()
    export_report(args.file, args.format, args.domain, args.output)
